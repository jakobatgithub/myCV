
import json
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# === Load JSON files ===
base_path = Path("data")
cv_data = json.load(open(base_path / "cv.json"))
skills_data = json.load(open(base_path / "skills.json"))
publications_data = json.load(open(base_path / "publications.json"))
thesis_data = json.load(open(base_path / "thesis.json"))
presentations_data = json.load(open(base_path / "presentations.json"))
posters_data = json.load(open(base_path / "posters.json"))

# === Helper Functions ===

def create_hyperlink(paragraph, text, url):
    if not url:
        paragraph.add_run(text)
        return paragraph
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def score_to_bar(score, max_score=10):
    return "■" * score + "□" * (max_score - score)

def format_paragraph(p, bold=False, italic=False, space_after=Pt(6), keep_with_next=False):
    run = p.runs[0] if p.runs else p.add_run()
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = space_after
    p.paragraph_format.keep_with_next = keep_with_next

# === Create Document ===
doc = Document()

# === Define Styles ===
styles = doc.styles
if "TitleStyle" not in styles:
    title_style = styles.add_style("TitleStyle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(20)
    title_style.font.bold = True

if "Heading1Style" not in styles:
    h1_style = styles.add_style("Heading1Style", WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = "Calibri"
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True

if "Heading2Style" not in styles:
    h2_style = styles.add_style("Heading2Style", WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = "Calibri"
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True

if "Heading3Style" not in styles:
    h3_style = styles.add_style("Heading3Style", WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = "Calibri"
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True

def styled_heading(text, level=1):
    style = "Heading1Style" if level == 1 else "Heading2Style"
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.keep_with_next = True

# === Title ===
doc.add_paragraph("Jakob Löber — Curriculum Vitae", style="TitleStyle")
doc.add_paragraph("PhD in Theoretical Physics | Data Science | Web Development | Augmented Reality", style="Heading3Style")

# === Professional Background ===
styled_heading("Professional Background", level=1)
for job in cv_data["professional_background"]:
    p1 = doc.add_paragraph(job["period"])
    format_paragraph(p1, bold=True, keep_with_next=True)

    p2 = doc.add_paragraph()
    p2.add_run(f"{job['position']} — ").italic = True
    if "link" in job:
        create_hyperlink(p2, job["company"], job["link"])
    else:
        p2.add_run(job["company"]).italic = True

    for detail in job.get("details", []):
        doc.add_paragraph(detail, style='List Bullet')

# === Education ===
styled_heading("Education", level=1)
for edu in cv_data["education"]:
    p1 = doc.add_paragraph(edu["period"])
    format_paragraph(p1, bold=True, keep_with_next=True)

    field = f" in {edu['field']}" if "field" in edu else ""
    p2 = doc.add_paragraph(f"{edu['degree']}{field} — {edu['institution']}")
    format_paragraph(p2, italic=True)

    if "grade" in edu:
        doc.add_paragraph(f"Grade: {edu['grade']}")
    if "focus" in edu:
        doc.add_paragraph(f"Focus: {', '.join(edu['focus'])}")
    if "title" in edu:
        doc.add_paragraph(f"Title: {edu['title']}")

# === Other Activities ===
styled_heading("Other Activities and Experiences", level=1)
for act in cv_data["other_activities_and_experiences"]:
    p1 = doc.add_paragraph(act["period"])
    format_paragraph(p1, bold=True, keep_with_next=True)

    p2 = doc.add_paragraph(f"{act['activity']} — {act['organization']}")
    format_paragraph(p2, italic=True)

    if "topics" in act:
        doc.add_paragraph(f"Topics: {', '.join(act['topics'])}")
    for detail in act.get("details", []):
        doc.add_paragraph(detail, style='List Bullet')

doc.add_page_break()

# === Skills ===
styled_heading("Skills Overview", level=1)

for category, skills in skills_data.items():
    doc.add_paragraph(category, style="Heading2Style")

    table = doc.add_table(rows=0, cols=4)

    # Iterate through skills in pairs of two
    for i in range(0, len(skills), 2):
        row = table.add_row().cells
        skill1 = skills[i]
        row[0].text = skill1["name"]
        row[1].text = score_to_bar(skill1["score"])

        if i + 1 < len(skills):
            skill2 = skills[i + 1]
            row[2].text = skill2["name"]
            row[3].text = score_to_bar(skill2["score"])
        else:
            # Leave the second pair empty if odd number of skills
            row[2].text = ""
            row[3].text = ""

    doc.add_paragraph()

# === Publications ===
styled_heading(f"Publications ({len(publications_data)})", level=1)
for i, pub in enumerate(publications_data, 1):
    authors = pub["authors"]
    title = f"“{pub['title']}”"
    journal = pub["journal"]
    p = doc.add_paragraph(f"{i}. {authors}, {title}, {journal}.")
    create_hyperlink(p, "", pub["url"])

doc.add_page_break()

# === Theses ===
styled_heading("Academic Theses", level=1)
dt = thesis_data["doctoral_thesis"]
styled_heading("Doctoral Thesis", level=2)
create_hyperlink(doc.add_paragraph("Title: "), dt["title"], dt["pdf_link"])
doc.add_paragraph("Supervisors: " + ", ".join(s["name"] for s in dt["supervisors"]))
doc.add_paragraph(f"Date of Defence: {dt['date_of_defence']}")
create_hyperlink(doc.add_paragraph("Defence Talk: "), dt["defence_talk"]["title"], dt["defence_talk"]["pdf_link"])
doc.add_paragraph("Abstract: " + dt["abstract"])

dip = thesis_data["diploma_thesis"]
styled_heading("Diploma Thesis", level=2)
create_hyperlink(doc.add_paragraph("Title: "), dip["title"], dip["pdf_link"])
doc.add_paragraph("Supervisors: " + ", ".join(s["name"] for s in dip["supervisors"]))
doc.add_paragraph("Abstract: " + dip["abstract"])

# === Presentations ===
styled_heading(f"Scientific Presentations ({len(presentations_data)})", level=1)
for t in presentations_data:
    p = doc.add_paragraph(f"{t['number']}. ")
    create_hyperlink(p, t["title"], t["link"])
    p.add_run(f" — {t['event']}")

# === Posters ===
styled_heading(f"Posters ({len(posters_data)})", level=1)
for p in posters_data:
    para = doc.add_paragraph(f"{p['number']}. ")
    create_hyperlink(para, p["title"], p.get("link"))
    para.add_run(f" — {p['event']}")

# === Save DOCX ===
doc.save("full_cv.docx")
print("✅ full_cv.docx created.")
